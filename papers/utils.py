from google import genai
from decouple import config
import json
import time
import os

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document

client = genai.Client(api_key=config('GEMINI_API_KEY'))

FONT_PATH = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'fonts', 'NotoSansTamil-Regular.ttf')
pdfmetrics.registerFont(TTFont('NotoTamil', FONT_PATH))


def generate_question_paper(subject, topics, difficulty):
    prompt = f"""
You are an expert exam question paper setter for a college internal assessment exam.

Generate a question paper for:
- Subject: {subject}
- Topics: {topics}
- Difficulty: {difficulty}

The paper MUST follow this EXACT structure:

PART A: 10 multiple-choice questions, each worth 1 mark.
Each question must have exactly 4 options (A, B, C, D) with one correct answer.
Every option MUST contain real, non-empty text - never leave any option blank.

PART B: 4 short-answer questions, each worth 5 marks (student answers any 2).

PART C: 3 topics, each worth 10 marks, with an EITHER/OR choice (2 alternative questions per topic).

Respond ONLY in valid JSON, with NO extra text, NO markdown fences.
Use this EXACT structure:

{{
  "part_a": [
    {{"q_no": 1, "question": "question text", "options": {{"A": "text", "B": "text", "C": "text", "D": "text"}}, "answer": "A"}}
  ],
  "part_b": [
    {{"q_no": 1, "question": "question text", "answer": "answer text"}}
  ],
  "part_c": [
    {{"q_no": 1, "question_a": "text", "question_b": "text", "answer_a": "text", "answer_b": "text"}}
  ]
}}

part_a must have exactly 10 items with 4 non-empty options each.
part_b must have exactly 4 items.
part_c must have exactly 3 items.
"""

    max_retries = 3
    last_error = None

    for attempt in range(max_retries):
        try:
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=prompt
            )

            raw_text = response.text.strip()
            raw_text = raw_text.replace("```json", "").replace("```", "").strip()

            data = json.loads(raw_text)
            return data

        except Exception as e:
            last_error = e
            print(f"Attempt {attempt + 1} failed: {e}")
            time.sleep(2)

    raise Exception(f"Failed after {max_retries} attempts. Last error: {last_error}")


def split_text(text, max_chars):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        if len(current) + len(word) + 1 <= max_chars:
            current += (" " if current else "") + word
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def format_options(options):
    if not options:
        return ""
    parts = []
    for key, value in options.items():
        value = (value or "").strip()
        if value:
            parts.append(f"({key}) {value}")
    return "   ".join(parts)


def draw_header(c, width, height, header_info):
    y = height - 50

    college_name = header_info.get('college_name') or ''
    course_name = header_info.get('course_name') or ''
    exam_type = header_info.get('exam_type') or ''
    subject_code = header_info.get('subject_code') or ''
    subject = header_info.get('subject') or ''
    semester = header_info.get('semester') or ''
    exam_date = header_info.get('exam_date') or ''

    c.setFont("NotoTamil", 15)
    c.drawCentredString(width / 2, y, college_name)
    y -= 22

    c.setFont("NotoTamil", 10)
    if course_name:
        c.drawCentredString(width / 2, y, course_name)
        y -= 16

    c.setFont("NotoTamil", 11)
    if exam_type:
        c.drawCentredString(width / 2, y, exam_type)
        y -= 18

    c.setFont("NotoTamil", 10)
    c.drawCentredString(width / 2, y, f"{subject_code}  -  {subject}")
    y -= 16

    c.drawCentredString(width / 2, y, f"{semester}      Date: {exam_date}      Maximum Marks: 50")
    y -= 20

    c.line(40, y, width - 40, y)
    y -= 20
    return y


def check_page_break(c, y, width, height, header_info, min_space=60):
    if y < min_space:
        c.showPage()
        y = draw_header(c, width, height, header_info)
    return y


def draw_wrapped(c, x, y, text, font="NotoTamil", size=10, max_chars=95, line_gap=14):
    c.setFont(font, size)
    lines = split_text(text, max_chars)
    for line in lines:
        c.drawString(x, y, line)
        y -= line_gap
    return y


def generate_pdf(data, header_info, filepath):
    c = canvas.Canvas(filepath, pagesize=A4)
    width, height = A4

    y = draw_header(c, width, height, header_info)

    c.setFont("NotoTamil", 11)
    c.drawString(40, y, "PART - A")
    c.drawRightString(width - 40, y, "(10 x 1 = 10 Marks)")
    y -= 10
    c.line(40, y, width - 40, y)
    y -= 18

    for q in data.get("part_a", []):
        y = check_page_break(c, y, width, height, header_info)
        q_text = f"{q.get('q_no', '')}. {q.get('question', '')}"
        y = draw_wrapped(c, 40, y, q_text, size=10)

        opt_line = format_options(q.get("options", {}))
        if opt_line:
            y = check_page_break(c, y, width, height, header_info)
            y = draw_wrapped(c, 55, y, opt_line, size=10, max_chars=100)
        y -= 8

    y -= 10
    y = check_page_break(c, y, width, height, header_info)

    c.setFont("NotoTamil", 11)
    c.drawString(40, y, "PART - B")
    c.drawRightString(width - 40, y, "(Answer any TWO) (2 x 5 = 10 Marks)")
    y -= 10
    c.line(40, y, width - 40, y)
    y -= 18

    for q in data.get("part_b", []):
        y = check_page_break(c, y, width, height, header_info)
        q_text = f"{q.get('q_no', '')}. {q.get('question', '')}"
        y = draw_wrapped(c, 40, y, q_text, size=10)
        y -= 10

    y -= 10
    y = check_page_break(c, y, width, height, header_info)

    c.setFont("NotoTamil", 11)
    c.drawString(40, y, "PART - C")
    c.drawRightString(width - 40, y, "(3 x 10 = 30 Marks)")
    y -= 10
    c.line(40, y, width - 40, y)
    y -= 18

    for idx, q in enumerate(data.get("part_c", []), start=1):
        y = check_page_break(c, y, width, height, header_info)
        qa_text = f"{idx}. a) {q.get('question_a', '')}"
        y = draw_wrapped(c, 40, y, qa_text, size=10)
        y -= 4

        y = check_page_break(c, y, width, height, header_info)
        c.setFont("NotoTamil", 10)
        c.drawCentredString(width / 2, y, "(OR)")
        y -= 16

        y = check_page_break(c, y, width, height, header_info)
        qb_text = f"    b) {q.get('question_b', '')}"
        y = draw_wrapped(c, 40, y, qb_text, size=10)
        y -= 14

    c.save()


def generate_answer_pdf(data, header_info, filepath):
    c = canvas.Canvas(filepath, pagesize=A4)
    width, height = A4

    y = draw_header(c, width, height, header_info)
    c.setFont("NotoTamil", 13)
    c.drawCentredString(width / 2, y, "ANSWER KEY")
    y -= 25

    c.setFont("NotoTamil", 11)
    c.drawString(40, y, "PART - A")
    y -= 18

    for q in data.get("part_a", []):
        y = check_page_break(c, y, width, height, header_info)
        options = q.get("options", {})
        correct_key = q.get("answer", "")
        correct_text = options.get(correct_key, "") if isinstance(options, dict) else ""
        text = f"{q.get('q_no', '')}. Correct Answer: ({correct_key}) {correct_text}"
        y = draw_wrapped(c, 40, y, text, size=10)
        y -= 6

    y -= 10
    y = check_page_break(c, y, width, height, header_info)

    c.setFont("NotoTamil", 11)
    c.drawString(40, y, "PART - B")
    y -= 18

    for q in data.get("part_b", []):
        y = check_page_break(c, y, width, height, header_info)
        y = draw_wrapped(c, 40, y, f"{q.get('q_no', '')}. {q.get('question', '')}", size=10)
        y = check_page_break(c, y, width, height, header_info)
        y = draw_wrapped(c, 55, y, f"Answer: {q.get('answer', '')}", size=10)
        y -= 10

    y -= 10
    y = check_page_break(c, y, width, height, header_info)

    c.setFont("NotoTamil", 11)
    c.drawString(40, y, "PART - C")
    y -= 18

    for idx, q in enumerate(data.get("part_c", []), start=1):
        y = check_page_break(c, y, width, height, header_info)
        y = draw_wrapped(c, 40, y, f"{idx}. a) {q.get('question_a', '')}", size=10)
        y = check_page_break(c, y, width, height, header_info)
        y = draw_wrapped(c, 55, y, f"Answer: {q.get('answer_a', '')}", size=10)
        y -= 8

        y = check_page_break(c, y, width, height, header_info)
        y = draw_wrapped(c, 40, y, f"    b) {q.get('question_b', '')}", size=10)
        y = check_page_break(c, y, width, height, header_info)
        y = draw_wrapped(c, 55, y, f"Answer: {q.get('answer_b', '')}", size=10)
        y -= 14

    c.save()


def generate_docx(data, header_info, filepath):
    doc = Document()

    doc.add_heading(header_info.get('college_name', ''), level=1)
    if header_info.get('course_name'):
        doc.add_paragraph(header_info.get('course_name'))
    doc.add_paragraph(header_info.get('exam_type', ''))
    doc.add_paragraph(f"{header_info.get('subject_code', '')} - {header_info.get('subject', '')}")
    doc.add_paragraph(f"{header_info.get('semester', '')}    Date: {header_info.get('exam_date', '')}    Maximum Marks: 50")
    doc.add_paragraph("")

    doc.add_heading("PART - A (10 x 1 = 10 Marks)", level=2)
    for q in data.get("part_a", []):
        doc.add_paragraph(f"{q.get('q_no', '')}. {q.get('question', '')}")
        opt_line = format_options(q.get("options", {}))
        if opt_line:
            doc.add_paragraph(opt_line)

    doc.add_heading("PART - B (Answer any TWO) (2 x 5 = 10 Marks)", level=2)
    for q in data.get("part_b", []):
        doc.add_paragraph(f"{q.get('q_no', '')}. {q.get('question', '')}")

    doc.add_heading("PART - C (3 x 10 = 30 Marks)", level=2)
    for idx, q in enumerate(data.get("part_c", []), start=1):
        doc.add_paragraph(f"{idx}. a) {q.get('question_a', '')}")
        doc.add_paragraph("(OR)")
        doc.add_paragraph(f"    b) {q.get('question_b', '')}")

    doc.save(filepath)