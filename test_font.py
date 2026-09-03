# --- PDF test (reportlab) ---
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Tamil font ah reportlab kitta "register" pannurom - idhu illama Tamil boxes (□□□) ah kaatum
pdfmetrics.registerFont(TTFont('TamilFont', 'fonts/Noto Sans Tamil Regular.ttf'))

c = canvas.Canvas("test_tamil.pdf")
c.setFont('TamilFont', 16)
c.drawString(100, 750, "வணக்கம் - இது தமிழ் எழுத்துரு சோதனை")
c.save()
print("PDF created: test_tamil.pdf")

# --- Word test (python-docx) ---
from docx import Document
from docx.oxml.ns import qn

doc = Document()
para = doc.add_paragraph()
run = para.add_run("வணக்கம் - இது தமிழ் எழுத்துரு சோதனை")
run.font.name = 'Nirmala UI'
# idhu Word ku sollrom "Tamil characters ku idha font use pannu" nu
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Nirmala UI')

doc.save("test_tamil.docx")
print("Word file created: test_tamil.docx")