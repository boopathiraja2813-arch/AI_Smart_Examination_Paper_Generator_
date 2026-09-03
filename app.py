import os
import io
import json
import re
import base64
import uuid
from datetime import date
from pathlib import Path

import streamlit as st
from dotenv import load_dotenv
from google import genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

load_dotenv()

# Each browser connection gets its own Streamlit session. Keep generated files
# in a per-session directory so simultaneous users can never overwrite each
# other's question papers or answer keys.
if "session_id" not in st.session_state:
    st.session_state["session_id"] = uuid.uuid4().hex
SESSION_ID = st.session_state["session_id"]

st.set_page_config(
    page_title="Artificial Intelligence Based Smart Examination Paper Generator",
    page_icon="📝",
    layout="wide",
)

st.markdown("""
<style>

/* Hide Streamlit top-right toolbar */
[data-testid="stToolbar"] {
    visibility: hidden !important;
    display: none !important;
}

/* Hide Deploy button */
[data-testid="stDeployButton"] {
    visibility: hidden !important;
    display: none !important;
}

/* Hide three-dot menu / header menu */
#MainMenu {
    visibility: hidden !important;
    display: none !important;
}

/* Hide header decoration */
[data-testid="stDecoration"] {
    display: none !important;
}

/* Hide status widget */
[data-testid="stStatusWidget"] {
    display: none !important;
}

</style>
""", unsafe_allow_html=True)

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "").strip()
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-3.5-flash").strip()
@st.cache_resource
def get_gemini_client():
    if not GEMINI_API_KEY:
        raise RuntimeError(
            "GEMINI_API_KEY is not configured. Add it to the project's .env file."
        )
    return genai.Client(api_key=GEMINI_API_KEY)

# Normalize the common typo. Gemini 3.5 Flash model id is exactly
# `gemini-3.5-flash` (no trailing s).
if GEMINI_MODEL == "gemini-3.5-flashs":
    GEMINI_MODEL = "gemini-3.5-flash"



# ============================================================
# GENERAL HELPERS
# ============================================================

def clean_json_text(text):
    if not text:
        return ""
    text = text.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.I)
    text = re.sub(r"\s*```$", "", text)
    return text.strip()


def extract_docx_text(uploaded_file):
    """Read a DOCX study material file into plain text."""
    doc = Document(io.BytesIO(uploaded_file.getvalue()))
    chunks = []

    for p in doc.paragraphs:
        if p.text.strip():
            chunks.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                chunks.append(row_text)

    return "\n".join(chunks)


def uploaded_material_part(uploaded_file):
    """Return a Gemini content part for PDF, or extracted text for DOCX."""
    name = uploaded_file.name.lower()
    data = uploaded_file.getvalue()

    if name.endswith(".pdf"):
        if len(data) > 25 * 1024 * 1024:
            raise ValueError("PDF is larger than 25 MB. Please upload a smaller unit PDF.")
        return {
            "inline_data": {
                "mime_type": "application/pdf",
                "data": base64.b64encode(data).decode("utf-8"),
            }
        }

    if name.endswith(".docx"):
        text = extract_docx_text(uploaded_file)
        if not text.strip():
            raise ValueError(f"No readable text found in {uploaded_file.name}.")
        return f"\n--- STUDY MATERIAL: {uploaded_file.name} ---\n{text}"

    raise ValueError("Only PDF and DOCX study materials are supported.")


def calculate_totals(meta):
    return sum(int(x['questions_to_attempt'])*int(x['marks_per_question']) for x in meta.get('sections', []))


def validate_pattern(meta):
    errors=[]; total=0
    for sec in meta.get('sections',[]):
        name=sec['name']; gen=int(sec['questions_to_generate']); att=int(sec['questions_to_attempt']); marks=int(sec['marks_per_question'])
        if gen<1: errors.append(f'{name}: questions to generate must be at least 1.')
        if att<1: errors.append(f'{name}: questions to attempt must be at least 1.')
        if att>gen: errors.append(f'{name}: attempt count cannot exceed generated count.')
        if marks<1: errors.append(f'{name}: marks per question must be at least 1.')
        total += att*marks
    if not meta.get('sections'): errors.append('Add at least one Part / Section.')
    if total != int(meta['total_marks']): errors.append(f'Total marks mismatch: configured sections = {total}, but Total Marks is {meta["total_marks"]}.')
    return errors,total


# ============================================================
# GEMINI QUESTION GENERATION

# ============================================================
# GEMINI QUESTION GENERATION
# ============================================================

def generate_question_paper(meta, uploaded_files):
    if not GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY is not configured. Add it to the project's .env file and restart the app.")
    client = get_gemini_client()
    material_parts=[uploaded_material_part(f) for f in uploaded_files]
    lines=[]
    for i,sec in enumerate(meta['sections']):
        lines.append(f"{i+1}. {sec['name']} | TYPE={sec['question_type']} | GENERATE={sec['questions_to_generate']} | ATTEMPT={sec['questions_to_attempt']} | MARKS={sec['marks_per_question']} | CHOICE={sec['choice_type']} | OR_CHOICE={bool(sec.get('or_choice',False))}")
    prompt="""
You are a strict college examination paper setter.
Use ONLY the uploaded study material. Never invent facts or answers. Never copy an existing question paper.
Follow the configured sections EXACTLY. Do not add missing sections, 5-mark questions, 10-mark questions, or extra patterns.

EXAM DETAILS:
College: {college}
Course: {course}
Subject: {subject}
Subject Code: {code}
Semester: {semester}
Set: {setname}
Maximum Marks: {marks}
Language: {lang}

CONFIGURED SECTIONS:
{sections}

Question types:
MCQ = four options A-D, one correct answer.
Paragraph / Descriptive = descriptive question and answer.
When OR_CHOICE=True for a Paragraph / Descriptive section, create TWO alternative descriptive questions for every configured question number. Put the second alternative in an `or_question` object with the same fields as the main question. The two alternatives are separated in the final paper by a centered `OR`. When OR_CHOICE=False, do not create `or_question`.
Only MCQ and Paragraph / Descriptive question types are allowed.

Language: English only, Tamil only, or English followed by Tamil according to the selected mode. Preserve the same numbering in bilingual output.
Generate EXACTLY the configured number of questions in each section. Number continuously from 1 across sections.

Return ONLY JSON:
{{
  "header_translations": {{"college_name":"","course_name":"","location":"","exam_session":"","subject_name":"","semester":"","set_name":""}},
  "sections": [
    {{"name":"Part A","question_type":"MCQ","instruction":"Answer all questions","passage":"","questions":[{{"q_no":1,"question":"","question_ta":"","options":{{"A":"","B":"","C":"","D":""}},"options_ta":{{"A":"","B":"","C":"","D":""}},"answer":"A","answer_text":"","answer_text_ta":"","co":"CO1","kl":"L1","or_question":null}}]}}
  ]
}}
For non-MCQ questions use answer_text/answer_text_ta. For Passage Based include section-level passage. For Match the Following use answer_text. Do not add fields or sections unnecessarily.
""".format(college=meta['college_name'],course=meta['course_name'],subject=meta['subject_name'],code=meta['subject_code'],semester=meta['semester'],setname=meta['set_name'],marks=meta['total_marks'],lang=meta['language_mode'],sections='\n'.join(lines))
    response = client.models.generate_content(
        model=GEMINI_MODEL,
        contents=material_parts + [prompt],
        config={
            "response_mime_type": "application/json"
        }
    )

    text = clean_json_text(getattr(response, "text", "") or "")

    try:
        data = json.loads(text)
    except Exception as exc:
        raise ValueError(
            "Gemini returned invalid JSON. Please try Generate again."
        ) from exc

    if len(data.get("sections", [])) != len(meta["sections"]):
        raise ValueError("AI returned an incorrect number of sections.")

    for got, want in zip(data["sections"], meta["sections"]):
        if len(got.get("questions", [])) != int(want["questions_to_generate"]):
            raise ValueError(
                f"AI generated an incorrect number of questions for {want['name']}."
            )

        got["name"] = want["name"]
        got["question_type"] = want["question_type"]
        got["marks_per_question"] = want["marks_per_question"]
        got["questions_to_attempt"] = want["questions_to_attempt"]
        got["choice_type"] = want["choice_type"]
        got["or_choice"] = bool(want.get("or_choice", False))
        got["passage_questions"] = want.get("passage_questions", 5)

        if got["or_choice"]:
            if got["question_type"] != "Paragraph / Descriptive":
                raise ValueError(
                    "Either / OR Choice is allowed only for Paragraph / Descriptive sections."
                )

            for qq in got.get("questions", []):
                if not isinstance(qq.get("or_question"), dict):
                    raise ValueError(
                        f"AI did not return the OR alternative for question {qq.get('q_no', '')}."
                    )
        else:
            for qq in got.get("questions", []):
                qq.pop("or_question", None)

    return data
    try: data=json.loads(text)
    except Exception as exc: raise ValueError('Gemini returned invalid JSON. Please try Generate again.') from exc
    if len(data.get('sections',[])) != len(meta['sections']): raise ValueError('AI returned an incorrect number of sections.')
    for got,want in zip(data['sections'],meta['sections']):
        if len(got.get('questions',[])) != int(want['questions_to_generate']): raise ValueError(f"AI generated an incorrect number of questions for {want['name']}.")
        got['name']=want['name']; got['question_type']=want['question_type']; got['marks_per_question']=want['marks_per_question']; got['questions_to_attempt']=want['questions_to_attempt']; got['choice_type']=want['choice_type']; got['or_choice']=bool(want.get('or_choice',False)); got['passage_questions']=want.get('passage_questions',5)
        if got['or_choice']:
            if got['question_type'] != 'Paragraph / Descriptive':
                raise ValueError('Either / OR Choice is allowed only for Paragraph / Descriptive sections.')
            for qq in got.get('questions', []):
                if not isinstance(qq.get('or_question'), dict):
                    raise ValueError(f"AI did not return the OR alternative for question {qq.get('q_no','')}.")
        else:
            for qq in got.get('questions', []):
                qq.pop('or_question', None)
    return data



# DOCX FORMATTING HELPERS
# ============================================================

def set_cell_shading(cell, fill="FFFFFF"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=35, start=55, bottom=35, end=55):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, size="6", color="000000", inside=True):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    edges = ["top", "left", "bottom", "right"]
    if inside:
        edges += ["insideH", "insideV"]
    for edge in edges:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_fixed_table_layout(table):
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    table.autofit = False


def _set_col_widths(table, widths):
    table.autofit = False
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(width * 1440)))
        grid.append(gc)
    for row in table.rows:
        for i, width in enumerate(widths):
            if i >= len(row.cells):
                continue
            cell = row.cells[i]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def _prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    tr_pr.append(cant)


def _set_cell_text(cell, text, bold=False, size=9.3, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    return p


def _set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, edge_data in kwargs.items():
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def set_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_doc_defaults(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.32)
    section.bottom_margin = Inches(0.36)
    section.left_margin = Inches(0.34)
    section.right_margin = Inches(0.34)
    section.header_distance = Inches(0.15)
    section.footer_distance = Inches(0.16)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(9.2)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    r = footer.add_run("Page ")
    r.font.name = "Times New Roman"
    r.font.size = Pt(7.5)
    set_page_number(footer)


def _new_cell_paragraph(cell, text="", bold=False, size=9.2, align=WD_ALIGN_PARAGRAPH.LEFT, after=0):
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = align
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    return p


def add_register_header(doc):
    """Reference-model register row: label followed by 14 boxes, positioned at the top-right."""
    table = doc.add_table(rows=1, cols=15)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    set_fixed_table_layout(table)
    set_table_borders(table, size="6")
    # Keep the label on the left of the 14 boxes, with the complete register row flush-right.
    _set_col_widths(table, [0.58] + [0.205] * 14)

    label = table.cell(0, 0)
    set_cell_margins(label, top=12, start=8, bottom=12, end=8)
    _set_cell_text(label, "Register\nNo:", bold=True, size=7.7, align=WD_ALIGN_PARAGRAPH.CENTER)
    label.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for i in range(1, 15):
        c = table.cell(0, i)
        set_cell_margins(c, top=12, start=2, bottom=12, end=2)
        _set_cell_text(c, "", size=7.7, align=WD_ALIGN_PARAGRAPH.CENTER)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)


def _centered_paragraph(doc, text, size, bold=True, after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    return p


def _lang_text(meta, key, english):
    mode = meta.get("language_mode", "English")
    tr = meta.get("header_translations", {}) or {}
    tamil = tr.get(key, "") or english
    if mode == "Tamil":
        return tamil
    if mode == "English + Tamil":
        return f"{english}\n{tamil}"
    return english


def _label_text(meta, english, tamil):
    mode = meta.get("language_mode", "English")
    if mode == "Tamil":
        return tamil
    if mode == "English + Tamil":
        return f"{english} / {tamil}"
    return english


def add_header(doc, meta):
    add_register_header(doc)
    _centered_paragraph(doc, _lang_text(meta, "college_name", meta["college_name"]).upper(), 13.2, True, 0)
    if meta.get("language_mode") == "Tamil":
        _centered_paragraph(doc, _lang_text(meta, "location", meta.get("college_location", "")).upper(), 8.2, True, 0)
    elif meta.get("language_mode") == "English + Tamil":
        _centered_paragraph(doc, meta.get("college_location", "").upper(), 8.0, True, 0)
        _centered_paragraph(doc, _lang_text(meta, "location", meta.get("college_location", "")).upper(), 8.0, True, 0)
    else:
        _centered_paragraph(doc, meta.get("college_location", "").upper(), 8.2, True, 0)
    _centered_paragraph(doc, "Accredited by NACC, Approved by the Govt. of Tamil Nadu and Affiliated with Periyar University, Salem.", 7.8, False, 0)
    _centered_paragraph(doc, _lang_text(meta, "exam_session", meta["exam_session"]).upper(), 10.0, True, 0)
    _centered_paragraph(doc, _lang_text(meta, "course_name", meta["course_name"]).upper(), 10.0, True, 0)
    _centered_paragraph(doc, _lang_text(meta, "semester", meta["semester"]).upper(), 9.8, True, 0)
    _centered_paragraph(doc, f"{meta['subject_code']}–{_lang_text(meta, 'subject_name', meta['subject_name'])}", 10.0, True, 0)
    _centered_paragraph(doc, _lang_text(meta, "set_name", meta["set_name"]).replace("SET-", "SET – "), 9.8, True, 3)

    dt = doc.add_table(rows=1, cols=3)
    dt.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_table_layout(dt)
    set_table_borders(dt, size="6")
    _set_col_widths(dt, [2.40, 2.45, 2.40])
    for c in dt.rows[0].cells:
        set_cell_margins(c, top=22, start=35, bottom=22, end=35)
    _set_cell_text(dt.cell(0, 0), f"{_label_text(meta, 'Date-', 'தேதி-')}   {meta['exam_date']}", bold=True, size=9.2, align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell_text(dt.cell(0, 1), f"{_label_text(meta, 'Time', 'நேரம்')}: {meta['exam_time']}", bold=True, size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(dt.cell(0, 2), f"{_label_text(meta, 'Maximum Marks', 'மொத்த மதிப்பெண்கள்')}: {meta['total_marks']}", bold=True, size=9.2, align=WD_ALIGN_PARAGRAPH.RIGHT)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)

def _add_part_heading(table, part_label, instruction, marks_text, right1="CO", right2="KL", meta=None):
    row = table.add_row()
    _prevent_row_split(row)
    merged = row.cells[0].merge(row.cells[1])
    for c in row.cells:
        set_cell_margins(c, top=22, start=25, bottom=22, end=25)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(c, "FFFFFF")

    p = merged.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_with_next = True
    r = p.add_run(part_label)
    r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(10.0)
    p2 = merged.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.0
    p2.paragraph_format.keep_with_next = True
    r2 = p2.add_run(f"{instruction} {marks_text}")
    r2.bold = True; r2.font.name = "Times New Roman"; r2.font.size = Pt(9.2)
    _set_cell_text(row.cells[2], right1, bold=True, size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(row.cells[3], right2, bold=True, size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
    return row

def _make_question_table(doc):
    table = doc.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_table_layout(table)
    set_table_borders(table, size="6")
    _set_col_widths(table, [0.48, 5.70, 0.57, 0.57])
    return table


def _add_option_grid(cell, options, size=9.0):
    opt = cell.add_table(rows=2, cols=2)
    opt.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_fixed_table_layout(opt)
    set_table_borders(opt, size="5")
    _set_col_widths(opt, [2.82, 2.82])
    vals = [
        (f"A", options["A"], f"C", options["C"]),
        (f"B", options["B"], f"D", options["D"]),
    ]
    for rr in range(2):
        for cc in range(2):
            c = opt.cell(rr, cc)
            set_cell_margins(c, top=10 if size >= 9.2 else 8, start=18, bottom=10 if size >= 9.2 else 8, end=18)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(f"{vals[rr][cc*2]}) {vals[rr][cc*2+1]}")
            r.font.name = "Times New Roman"; r.font.size = Pt(size)
    return opt


def _allow_row_split(row):
    """Explicitly allow a question row to continue onto the next page.

    This is intentional for long question/option rows: Word otherwise moves the
    entire row to the next page and leaves a large blank area at the bottom of
    the current page.
    """
    tr_pr = row._tr.get_or_add_trPr()
    cant = tr_pr.find(qn("w:cantSplit"))
    if cant is not None:
        tr_pr.remove(cant)


def add_question_table_row(table, q_no, question, options=None, co="", kl="", size=9.2, compact=False):
    row = table.add_row()
    # Keep each question together for clean student printing. Descriptive/Part-B
    # rows use a tighter, print-friendly height so later parts can use the same
    # page instead of being pushed to an unnecessary new page. MCQ rows keep the
    # original college-template spacing.
    _prevent_row_split(row)
    if options:
        row.height = Inches(0.75)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    elif compact:
        row.height = Inches(0.40)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    elif size >= 9.2:
        row.height = Inches(0.58)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    cells = row.cells
    vertical_pad = 10 if compact else (30 if options else 26)
    for c in cells:
        set_cell_margins(c, top=vertical_pad, start=20, bottom=vertical_pad, end=20)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    _set_cell_text(cells[0], str(q_no), bold=False, size=9.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(cells[1], question, size=size, align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell_text(cells[2], co, size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(cells[3], kl, size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)

    if options:
        _add_option_grid(cells[1], options, size=size - 0.2)
    return row


def add_part_heading_table(doc, part_label, instruction, marks_text, right1="CO", right2="KL", meta=None):
    table = _make_question_table(doc)
    _add_part_heading(table, part_label, instruction, marks_text, right1, right2, meta)
    return table


def add_answer_key_header(doc, meta):
    """Teacher-facing answer-key header: no register box or location block."""
    _centered_paragraph(doc, _lang_text(meta, "college_name", meta["college_name"]).upper(), 13.0, True, 1)
    _centered_paragraph(doc, f"{meta['subject_code']} – {_lang_text(meta, 'subject_name', meta['subject_name'])}".upper(), 10.5, True, 0)
    _centered_paragraph(doc, _lang_text(meta, "set_name", meta["set_name"]).replace("SET-", "SET – "), 10.0, True, 4)

def add_signatory_footer(doc):
    """Exact model-style approval row: one bordered row around all five labels."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SIGNATURE / APPROVAL")
    r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(9.0)

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_table_layout(table)
    set_table_borders(table, size="6")
    _set_col_widths(table, [1.46] * 5)
    names = ["SUBJECT IN-CHARGE", "HOD", "PH", "EXAM CELL", "PRINCIPAL"]
    for i, name in enumerate(names):
        c = table.cell(0, i)
        set_cell_margins(c, top=18, start=10, bottom=18, end=10)
        _set_cell_text(c, name, bold=True, size=8.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _bilingual(meta, english, tamil=None):
    mode = meta.get("language_mode", "English")
    tamil = tamil or ""
    if mode == "Tamil":
        return tamil or english
    if mode == "English + Tamil":
        return f"{english}\n{tamil or english}"
    return english


def _translated_question(q, key, meta):
    ta = q.get(f"{key}_ta", "")
    return _bilingual(meta, q.get(key, ""), ta)


def _translated_option(q, key, meta):
    opts = q.get("options", {}) or {}
    ta_opts = q.get("options_ta", {}) or {}
    return _bilingual(meta, opts.get(key, ""), ta_opts.get(key, ""))


def _add_part_table(doc, part_label, instruction, marks_text, questions, kind, meta):
    labels = {
        "A": ("PART-A", "பகுதி - A", f"Answer all questions", f"அனைத்து கேள்விகளுக்கும் பதிலளிக்கவும்"),
        "B": ("PART-B", "பகுதி - B", f"Answer any {meta['part_b_attempt']} questions", f"{meta['part_b_attempt']} கேள்விகளுக்கு மட்டும் பதிலளிக்கவும்"),
        "C": ("PART-C", "பகுதி - C", "Answer all questions", "அனைத்து கேள்விகளுக்கும் பதிலளிக்கவும்"),
    }[kind]
    mode = meta.get("language_mode", "English")
    if mode == "Tamil":
        part_label = labels[1]
        instruction = labels[3]
    elif mode == "English + Tamil":
        part_label = f"{labels[0]} / {labels[1]}"
        instruction = f"{labels[2]} / {labels[3]}"
    else:
        part_label = labels[0]
        instruction = labels[2]

    table = add_part_heading_table(
        doc, part_label, instruction, marks_text,
        right1="CLO" if kind == "B" else "CO", right2="KL", meta=meta
    )
    if kind == "A":
        for q in questions:
            q2 = dict(q)
            q2["question"] = _translated_question(q, "question", meta)
            opts = {k: _translated_option(q, k, meta) for k in ("A", "B", "C", "D")}
            add_question_table_row(table, q["q_no"], q2["question"], opts, q.get("co", "CO1"), q.get("kl", "L1"), size=9.4)
    elif kind == "B":
        for q in questions:
            add_question_table_row(table, q["q_no"], _translated_question(q, "question", meta), None, q.get("co", "CO1"), q.get("kl", "L2"), size=9.2)
    else:
        for q in questions:
            if meta["part_c_or"]:
                row_a = table.add_row(); _allow_row_split(row_a)
                row_or = table.add_row(); _allow_row_split(row_or)
                row_b = table.add_row(); _allow_row_split(row_b)
                num_cell = row_a.cells[0].merge(row_or.cells[0]).merge(row_b.cells[0])
                set_cell_margins(num_cell, top=12, start=18, bottom=12, end=18)
                _set_cell_text(num_cell, str(q["q_no"]), size=9.0, align=WD_ALIGN_PARAGRAPH.CENTER)
                qa = _translated_question(q, "question_a", meta)
                qb = _translated_question(q, "question_b", meta)
                for row, text, co, kl in ((row_a, f"A    {qa}", q.get("co", "CO1"), q.get("kl", "L3")),
                                          (row_b, f"B    {qb}", q.get("co", "CO1"), q.get("kl", "L3"))):
                    _set_cell_text(row.cells[1], text, size=9.1, align=WD_ALIGN_PARAGRAPH.LEFT)
                    _set_cell_text(row.cells[2], co, size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
                    _set_cell_text(row.cells[3], kl, size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
                    for c in row.cells:
                        set_cell_margins(c, top=14, start=22, bottom=14, end=22)
                or_text = "OR" if mode == "English" else ("அல்லது" if mode == "Tamil" else "OR / அல்லது")
                _set_cell_text(row_or.cells[1], or_text, bold=True, size=9.0, align=WD_ALIGN_PARAGRAPH.CENTER)
                _set_cell_text(row_or.cells[2], "", size=8)
                _set_cell_text(row_or.cells[3], "", size=8)
                for c in row_or.cells:
                    set_cell_margins(c, top=6, start=18, bottom=6, end=18)
            else:
                add_question_table_row(table, q["q_no"], _translated_question(q, "question_a", meta), None, q.get("co", "CO1"), q.get("kl", "L3"), size=9.1)
    return table

def generate_question_docx(meta, data, filepath):
    doc = Document()
    set_doc_defaults(doc)
    add_header(doc, meta)

    a_total, b_total, c_total, _ = calculate_totals(meta)

    # Flow all sections continuously. Do NOT force a page break between parts.
    # Word will naturally carry rows to the next page when the current page is full,
    # so if Part-A has 10 questions and there is room below it, Part-B starts there;
    # if Part-A has 15/20/etc. questions, the remaining questions continue at the
    # top of the next page without leaving an artificial blank area.
    _add_part_table(
        doc, "PART-A", "Answer all questions", f"({meta['part_a_count']} x {meta['part_a_mark']} = {a_total})",
        data["part_a"], "A", meta
    )
    _add_part_table(
        doc, "PART-B", f"Answer any {meta['part_b_attempt']} questions", f"({meta['part_b_mark']} x {meta['part_b_attempt']} = {b_total})",
        data["part_b"], "B", meta
    )
    _add_part_table(
        doc, "PART-C", "Answer all questions", f"({meta['part_c_count']} x {meta['part_c_mark']} = {c_total})",
        data["part_c"], "C", meta
    )
    add_signatory_footer(doc)
    doc.save(filepath)


def generate_answer_key_docx(meta, data, filepath):
    doc = Document()
    set_doc_defaults(doc)
    add_answer_key_header(doc, meta)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(_label_text(meta, "ANSWER KEY", "விடைக்குறிப்பு")); r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(13)

    def section_title(en, ta):
        mode = meta.get("language_mode", "English")
        if mode == "Tamil": return ta
        if mode == "English + Tamil": return f"{en} / {ta}"
        return en

    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = h.add_run(section_title("PART-A", "பகுதி - A")); rr.bold=True; rr.font.name="Times New Roman"; rr.font.size=Pt(10.5)
    for q in data["part_a"]:
        ans = q["answer"]
        eng = f"{q['q_no']}. ({ans}) {q['options'][ans]}"
        ta = f"{q['q_no']}. ({ans}) {(q.get('options_ta', {}) or {}).get(ans, q['options'][ans])}"
        text = _bilingual(meta, eng, ta)
        p = doc.add_paragraph(text); p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.12
        for run in p.runs: run.font.name="Times New Roman"; run.font.size=Pt(9.5)

    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = h.add_run(section_title("PART-B", "பகுதி - B")); rr.bold=True; rr.font.name="Times New Roman"; rr.font.size=Pt(10.5)
    for q in data["part_b"]:
        eng = f"{q['q_no']}. {q['answer']}"
        ta = f"{q['q_no']}. {q.get('answer_ta', q['answer'])}"
        p = doc.add_paragraph(_bilingual(meta, eng, ta)); p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(9); p.paragraph_format.line_spacing=1.12
        for run in p.runs: run.font.name="Times New Roman"; run.font.size=Pt(9.5)

    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = h.add_run(section_title("PART-C", "பகுதி - C")); rr.bold=True; rr.font.name="Times New Roman"; rr.font.size=Pt(10.5)
    for q in data["part_c"]:
        eng_a = f"{q['q_no']}. (A) {q['answer_a']}"; ta_a = f"{q['q_no']}. (A) {q.get('answer_a_ta', q['answer_a'])}"
        p = doc.add_paragraph(_bilingual(meta, eng_a, ta_a)); p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.12
        if meta["part_c_or"]:
            eng_b = f"{q['q_no']}. (B) {q['answer_b']}"; ta_b = f"{q['q_no']}. (B) {q.get('answer_b_ta', q['answer_b'])}"
            p = doc.add_paragraph(_bilingual(meta, eng_b, ta_b)); p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.12
        
    doc.save(filepath)

# ============================================================
def _section_title(sec, meta):
    en = sec.get("name", "Part")
    ta = sec.get("name_ta", en)
    if meta.get("language_mode") == "Tamil": return ta
    if meta.get("language_mode") == "English + Tamil": return f"{en} / {ta}"
    return en


def _section_instruction(sec, meta):
    att = int(sec.get("questions_to_attempt", 0))
    gen = int(sec.get("questions_to_generate", len(sec.get("questions", []))))
    choice = sec.get("choice_type", "Answer All")
    if choice == "Answer Any" and att < gen:
        return f"Answer any {att} questions"
    return "Answer all questions"


def _add_dynamic_section_table(doc, sec, meta, answer_key=False):
    """Render one configured section using the unchanged four-column college table.

    For the first MCQ section, page 1 is deliberately capped at 10 questions.
    The continuation starts at the top of page 2, and later sections are allowed
    to use whatever space remains on that page before Word naturally continues.
    """
    title = _section_title(sec, meta)
    instruction = _section_instruction(sec, meta)
    gen = int(sec.get("questions_to_generate", len(sec.get("questions", []))))
    att = int(sec.get("questions_to_attempt", gen))
    marks = int(sec.get("marks_per_question", 1))
    total = att * marks
    title = title.upper()
    if meta.get("language_mode") == "Tamil":
        title = sec.get("name_ta", title).upper()
        instruction = "அனைத்து கேள்விகளுக்கும் பதிலளிக்கவும்" if att == gen else f"{att} கேள்விகளுக்கு மட்டும் பதிலளிக்கவும்"
    elif meta.get("language_mode") == "English + Tamil":
        title = f"{title} / {sec.get('name_ta', title)}"
        instruction = f"{instruction} / {('அனைத்து கேள்விகளுக்கும் பதிலளிக்கவும்' if att == gen else f'{att} கேள்விகளுக்கு மட்டும் பதிலளிக்கவும்')}"

    questions = list(sec.get("questions", []))
    qtype = sec.get("question_type", "MCQ")
    # Part B is intentionally compact while retaining the same college table
    # and readable font. This lets Part C use available page space naturally.
    compact_rows = (title.replace("/", " ").strip().upper().startswith("PART B") and qtype == "Paragraph / Descriptive")

    # Only the first MCQ section is split at exactly 10 questions. This keeps
    # the first page clean while allowing the continuation page to flow directly
    # into Part B without an unnecessary page break.
    chunks = [questions]
    split_first_page = (not answer_key and qtype == "MCQ" and len(questions) > 10 and title.replace("/", " ").strip().upper().startswith("PART A"))
    if split_first_page:
        chunks = [questions[:10], questions[10:]]

    for chunk_index, chunk in enumerate(chunks):
        if chunk_index > 0:
            doc.add_page_break()
        table = _make_question_table(doc)
        # The section heading appears only on the first chunk. The continuation
        # page intentionally begins with the remaining question numbers, keeping
        # the exact college table style without inserting an extra title block.
        if chunk_index == 0:
            _add_part_heading(table, title, instruction, f"({gen} x {marks} = {total})", "CO", "KL", meta)

        for q in chunk:
            q_no = q.get("q_no", "")
            co = q.get("co", "CO1")
            kl = q.get("kl", "L1")
            if answer_key:
                if qtype == "MCQ":
                    ans = q.get("answer", "")
                    answer_text = q.get("options", {}).get(ans, ans)
                    add_question_table_row(table, q_no, f"({ans}) {answer_text}", None, co, kl, size=9.2)
                else:
                    answer_text = q.get("answer_text", q.get("answer", ""))
                    if qtype == "Paragraph / Descriptive" and sec.get("or_choice"):
                        answer_text = f"A) {answer_text}"
                    add_question_table_row(table, q_no, answer_text, None, co, kl, size=9.2, compact=compact_rows)
                    if qtype == "Paragraph / Descriptive" and sec.get("or_choice") and isinstance(q.get("or_question"), dict):
                        oq = q.get("or_question", {})
                        or_row = table.add_row(); _allow_row_split(or_row)
                        for c in or_row.cells:
                            set_cell_margins(c, top=6 if compact_rows else 10, start=20, bottom=6 if compact_rows else 10, end=20)
                        _set_cell_text(or_row.cells[0], "", size=9.0, align=WD_ALIGN_PARAGRAPH.CENTER)
                        _set_cell_text(or_row.cells[1], "OR", bold=True, size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
                        _set_cell_text(or_row.cells[2], "", size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
                        _set_cell_text(or_row.cells[3], "", size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
                        add_question_table_row(table, q_no, f"B) {oq.get('answer_text', oq.get('answer', ''))}", None, oq.get("co", co), oq.get("kl", kl), size=9.2, compact=compact_rows)
                continue

            question = q.get("question", "")
            if meta.get("language_mode") == "Tamil":
                question = q.get("question_ta", question)
            elif meta.get("language_mode") == "English + Tamil":
                question = f"{question}\n{q.get('question_ta', question)}"

            options = None
            if qtype == "MCQ":
                opts = q.get("options", {}) or {}
                ta_opts = q.get("options_ta", {}) or {}
                if meta.get("language_mode") == "Tamil":
                    opts = ta_opts or opts
                elif meta.get("language_mode") == "English + Tamil":
                    opts = {k: f"{opts.get(k, '')} / {ta_opts.get(k, opts.get(k, ''))}" for k in "ABCD"}
                options = {k: opts.get(k, "") for k in "ABCD"}

            if qtype == "Paragraph / Descriptive" and sec.get("or_choice"):
                question = f"A) {question}"
            add_question_table_row(table, q_no, question, options, co, kl, size=10.0, compact=compact_rows)

            if (not answer_key) and qtype == "Paragraph / Descriptive" and sec.get("or_choice") and isinstance(q.get("or_question"), dict):
                oq = q.get("or_question", {})
                oq_text = oq.get("question", "")
                if meta.get("language_mode") == "Tamil":
                    oq_text = oq.get("question_ta", oq_text)
                elif meta.get("language_mode") == "English + Tamil":
                    oq_text = f"{oq_text}\n{oq.get('question_ta', oq_text)}"
                or_row = table.add_row(); _allow_row_split(or_row)
                for c in or_row.cells:
                    set_cell_margins(c, top=6 if compact_rows else 10, start=20, bottom=6 if compact_rows else 10, end=20)
                _set_cell_text(or_row.cells[0], "", size=9.0, align=WD_ALIGN_PARAGRAPH.CENTER)
                _set_cell_text(or_row.cells[1], "OR", bold=True, size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
                _set_cell_text(or_row.cells[2], "", size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
                _set_cell_text(or_row.cells[3], "", size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)

                alt_row = table.add_row(); _allow_row_split(alt_row)
                for c in alt_row.cells:
                    set_cell_margins(c, top=8 if compact_rows else 14, start=20, bottom=8 if compact_rows else 14, end=20)
                    c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                _set_cell_text(alt_row.cells[0], "", size=9.0, align=WD_ALIGN_PARAGRAPH.CENTER)
                _set_cell_text(alt_row.cells[1], f"B) {oq_text}", size=9.2, align=WD_ALIGN_PARAGRAPH.LEFT)
                _set_cell_text(alt_row.cells[2], oq.get("co", co), size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
                _set_cell_text(alt_row.cells[3], oq.get("kl", kl), size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)

        passage = sec.get("passage")
        if qtype == "Passage Based" and passage and not answer_key:
            row = table.add_row(); _allow_row_split(row)
            merged = row.cells[0].merge(row.cells[1]).merge(row.cells[2]).merge(row.cells[3])
            _set_cell_text(merged, passage, bold=False, size=9.0, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_margins(merged, top=18, start=22, bottom=18, end=22)

    return table

def generate_question_docx(meta, data, filepath):
    doc = Document()
    set_doc_defaults(doc)
    add_header(doc, meta)

    # Print-friendly flow: Part A has a maximum of 10 MCQ rows on page 1.
    # Remaining Part A rows continue at the top of page 2, then Part B uses
    # any remaining space on that page before Word naturally continues later.
    for sec_index, sec in enumerate(data.get("sections", [])):
        qcount = len(sec.get("questions", []))
        qtype = sec.get("question_type", "MCQ")
        # Sections flow continuously. In particular, when Part A continues
        # onto page 2, Part B must use the remaining space on that same page
        # instead of being pushed to page 3 unnecessarily. Word will create a
        # new page only when the current page has no room.
        _add_dynamic_section_table(doc, sec, meta, False)
        if sec_index == 0 and qtype == "MCQ" and qcount > 10:
            # Split Part A after exactly ten questions. The remaining Part A
            # questions continue naturally on page 2, followed by Part B.
            # Rebuild is avoided here; the section table itself is allowed to
            # flow, so the explicit break is handled by the helper below.
            pass

    if meta.get("include_signature", False):
        add_signatory_footer(doc)
    doc.save(filepath)


def generate_answer_key_docx(meta, data, filepath):
    doc = Document()
    set_doc_defaults(doc)
    add_answer_key_header(doc, meta)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(_label_text(meta, "ANSWER KEY", "விடைக்குறிப்பு")); r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(13)
    for sec in data.get("sections", []):
        _add_dynamic_section_table(doc, sec, meta, True)
    doc.save(filepath)


# ============================================================
# FRONTEND — FINAL AI EXAM GENERATOR UI
# The frontend is intentionally separate from the DOCX template functions above.
# ============================================================
st.markdown("""
<style>
/* =========================================================
   COLLEGE EDITION — FRONTEND ONLY
   Backend, AI generation, validation and document formatting
   are intentionally untouched.
   ========================================================= */

:root {
    --college-navy: #102a43;
    --college-navy-2: #173f5f;
    --college-teal: #0f766e;
    --college-teal-2: #14b8a6;
    --college-blue: #2563eb;
    --college-gold: #d4a72c;
    --page-bg: #f4f7fb;
    --surface: #ffffff;
    --surface-soft: #f8fafc;
    --line: #d9e2ec;
    --ink: #102a43;
    --muted: #627d98;
}

/* ---------- Main application ---------- */
[data-testid="stAppViewContainer"] {
    background:
        radial-gradient(circle at 85% 5%, rgba(20,184,166,.09), transparent 28%),
        linear-gradient(180deg, #f8fafc 0%, #f4f7fb 100%);
}

[data-testid="stHeader"] {
    background: rgba(255,255,255,.94) !important;
    border-bottom: 1px solid rgba(217,226,236,.8);
}

/* Keep the Streamlit menu unobtrusive */
[data-testid="stToolbar"] {
    background: transparent !important;
}

/* ---------- Sidebar ---------- */
[data-testid="stSidebar"] {
    background:
        linear-gradient(180deg, #0b1f33 0%, #102a43 52%, #123d52 100%) !important;
    border-right: 1px solid rgba(255,255,255,.08);
}

[data-testid="stSidebar"] > div:first-child {
    padding-top: 1.1rem;
}

[data-testid="stSidebar"] * {
    color: #f7fbff !important;
}

/* Sidebar heading / brand */
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3 {
    color: #ffffff !important;
    letter-spacing: .1px;
}

[data-testid="stSidebar"] h2 {
    font-size: 1.28rem !important;
    font-weight: 800 !important;
    margin-bottom: .15rem !important;
}

[data-testid="stSidebar"] .stCaption,
[data-testid="stSidebar"] small {
    color: #a9c2d4 !important;
    letter-spacing: .8px;
    font-weight: 700;
}

/* Sidebar separators */
[data-testid="stSidebar"] hr {
    border-color: rgba(255,255,255,.13) !important;
    margin: 1rem 0 !important;
}

/* ---------- Sidebar labels ---------- */
[data-testid="stSidebar"] label,
[data-testid="stSidebar"] .stMarkdown p {
    color: #edf6fb !important;
}

[data-testid="stSidebar"] label {
    font-size: .82rem !important;
    font-weight: 700 !important;
}

/* ---------- Sidebar text inputs ---------- */
[data-testid="stSidebar"] input,
[data-testid="stSidebar"] textarea {
    background: #f7fafc !important;
    color: #16324a !important;
    border: 1px solid #b9cad8 !important;
    border-radius: 9px !important;
    min-height: 42px !important;
    font-weight: 600 !important;
    box-shadow: inset 0 1px 2px rgba(16,42,67,.04) !important;
}

[data-testid="stSidebar"] input::placeholder,
[data-testid="stSidebar"] textarea::placeholder {
    color: #829ab0 !important;
}

[data-testid="stSidebar"] input:focus,
[data-testid="stSidebar"] textarea:focus {
    border-color: var(--college-teal) !important;
    box-shadow: 0 0 0 2px rgba(20,184,166,.18) !important;
}

/* ---------- Sidebar select / dropdown controls ----------
   These selectors remove the unwanted white split/box effect
   while keeping the dropdown readable and professional.
   ---------- */
[data-testid="stSidebar"] [data-baseweb="select"] > div,
[data-testid="stSidebar"] [data-baseweb="input"] > div,
[data-testid="stSidebar"] [data-testid="stNumberInput"] {
    background: #f7fafc !important;
    color: #16324a !important;
    border: 1px solid #b9cad8 !important;
    border-radius: 9px !important;
    min-height: 42px !important;
    box-shadow: none !important;
}

[data-testid="stSidebar"] [data-baseweb="select"] span,
[data-testid="stSidebar"] [data-baseweb="select"] input,
[data-testid="stSidebar"] [data-baseweb="input"] input {
    color: #16324a !important;
    font-weight: 600 !important;
}

[data-testid="stSidebar"] [data-baseweb="select"] svg {
    fill: #315b76 !important;
}

[data-testid="stSidebar"] [data-testid="stNumberInput"] input {
    background: transparent !important;
    color: #16324a !important;
    border: 0 !important;
    font-weight: 700 !important;
}

[data-testid="stSidebar"] [data-testid="stNumberInput"] button {
    background: #e7eef5 !important;
    color: #173f5f !important;
    border-left: 1px solid #cbd8e4 !important;
}

[data-testid="stSidebar"] [data-testid="stNumberInput"] button:hover {
    background: #d7e6ef !important;
    color: #0f766e !important;
}

/* Dropdown popup */
[data-baseweb="popover"],
[data-baseweb="menu"] {
    border-radius: 10px !important;
}

[data-baseweb="menu"] {
    background: #ffffff !important;
    border: 1px solid #d9e2ec !important;
    box-shadow: 0 14px 35px rgba(16,42,67,.16) !important;
}

[data-baseweb="menu"] * {
    color: #16324a !important;
}

[data-baseweb="option"]:hover {
    background: #e9f7f5 !important;
}

/* ---------- Sidebar radio / checkbox ---------- */
[data-testid="stSidebar"] [role="radiogroup"] label,
[data-testid="stSidebar"] [data-testid="stCheckbox"] label {
    color: #eef7fb !important;
    font-weight: 600 !important;
}

[data-testid="stSidebar"] [data-testid="stCheckbox"] {
    margin-top: .25rem;
}

/* ---------- Sidebar buttons ---------- */
[data-testid="stSidebar"] button {
    border-radius: 9px !important;
    font-weight: 750 !important;
}

[data-testid="stSidebar"] button[kind="secondary"] {
    background: rgba(255,255,255,.09) !important;
    border: 1px solid rgba(255,255,255,.17) !important;
    color: #ffffff !important;
}

[data-testid="stSidebar"] button[kind="secondary"]:hover {
    background: rgba(20,184,166,.18) !important;
    border-color: rgba(20,184,166,.55) !important;
}

/* ---------- Main content width / spacing ---------- */
.block-container {
    max-width: 1180px !important;
    padding-top: 2.0rem !important;
    padding-bottom: 4rem !important;
}

/* ---------- Hero ---------- */
.ai-hero {
    position: relative;
    overflow: hidden;
    padding: 34px 38px 32px;
    margin: 0 0 24px;
    border-radius: 22px;
    color: #ffffff;
    background:
        linear-gradient(135deg, #102a43 0%, #173f5f 58%, #0f766e 100%);
    box-shadow: 0 18px 45px rgba(16,42,67,.18);
    border: 1px solid rgba(255,255,255,.12);
}

.ai-hero::after {
    content: "";
    position: absolute;
    width: 220px;
    height: 220px;
    right: -80px;
    top: -90px;
    border-radius: 50%;
    background: rgba(255,255,255,.08);
}

.ai-hero::before {
    content: "";
    position: absolute;
    width: 110px;
    height: 110px;
    right: 110px;
    bottom: -75px;
    border-radius: 50%;
    background: rgba(212,167,44,.13);
}

.ai-hero h1 {
    position: relative;
    z-index: 1;
    margin: 0 !important;
    max-width: 920px;
    font-size: clamp(28px, 3.1vw, 42px) !important;
    line-height: 1.14 !important;
    font-weight: 850 !important;
    letter-spacing: -.6px !important;
    color: #ffffff !important;
}

.ai-hero p {
    position: relative;
    z-index: 1;
    margin: 13px 0 0 !important;
    max-width: 860px;
    font-size: 16px !important;
    line-height: 1.55 !important;
    color: #dbeaf3 !important;
}

.ai-badge {
    position: relative;
    z-index: 1;
    display: inline-flex;
    align-items: center;
    margin-top: 18px;
    padding: 8px 13px;
    border-radius: 999px;
    background: rgba(255,255,255,.10);
    border: 1px solid rgba(255,255,255,.24);
    color: #ffffff !important;
    font-size: 12.5px;
    font-weight: 700;
}

/* ---------- Info messages ---------- */
div[data-testid="stAlert"] {
    border-radius: 12px !important;
    border: 1px solid #cfe0ef !important;
    box-shadow: 0 5px 18px rgba(16,42,67,.05) !important;
}

div[data-testid="stAlert"] p {
    color: #214b68 !important;
}

/* ---------- Content cards ---------- */
.ai-card {
    background: rgba(255,255,255,.97);
    border: 1px solid #dce6ee;
    border-radius: 18px;
    padding: 23px 25px;
    margin: 0 0 2px;
    box-shadow: 0 9px 28px rgba(16,42,67,.065);
}

.step-title {
    font-size: 23px;
    line-height: 1.25;
    font-weight: 850;
    color: var(--college-navy);
    margin-bottom: 5px;
}

.step-sub {
    color: var(--muted);
    font-size: 14.5px;
    line-height: 1.55;
    margin-bottom: 17px;
}

/* ---------- Main uploader ---------- */
[data-testid="stFileUploader"] {
    margin-top: 3px;
}

[data-testid="stFileUploader"] section {
    background: #f7fafc !important;
    border: 2px dashed #b8cad8 !important;
    border-radius: 13px !important;
    padding: 10px !important;
}

[data-testid="stFileUploader"] section:hover {
    border-color: var(--college-teal) !important;
    background: #f1faf9 !important;
}

[data-testid="stFileUploader"] button {
    background: #ffffff !important;
    color: #16324a !important;
    border: 1px solid #b9cad8 !important;
    border-radius: 8px !important;
    font-weight: 750 !important;
}

[data-testid="stFileUploader"] small,
[data-testid="stFileUploader"] span {
    color: #627d98 !important;
}

/* Uploaded file chips */
[data-testid="stFileUploaderFile"] {
    background: #eef7f5 !important;
    border: 1px solid #c6e5df !important;
    border-radius: 9px !important;
}

/* ---------- Main buttons ---------- */
div.stButton > button {
    border-radius: 10px !important;
    min-height: 44px !important;
    font-weight: 800 !important;
    letter-spacing: .1px;
    transition: transform .15s ease, box-shadow .15s ease, background .15s ease !important;
}

div.stButton > button:hover {
    transform: translateY(-1px);
}

div.stButton > button[kind="primary"] {
    background: linear-gradient(135deg, #0f766e, #0b8f86) !important;
    color: #ffffff !important;
    border: 0 !important;
    min-height: 52px !important;
    border-radius: 12px !important;
    font-weight: 850 !important;
    box-shadow: 0 9px 22px rgba(15,118,110,.22) !important;
}

div.stButton > button[kind="primary"]:hover {
    background: linear-gradient(135deg, #0b625c, #087a72) !important;
    box-shadow: 0 12px 27px rgba(15,118,110,.27) !important;
}

/* ---------- Expanders / preview ---------- */
[data-testid="stExpander"] {
    background: #ffffff !important;
    border: 1px solid #dce6ee !important;
    border-radius: 12px !important;
    margin-bottom: 9px !important;
    box-shadow: 0 4px 15px rgba(16,42,67,.045) !important;
}

[data-testid="stExpander"] summary {
    color: var(--college-navy) !important;
    font-weight: 800 !important;
}

/* ---------- Download buttons ---------- */
[data-testid="stDownloadButton"] button {
    background: #ffffff !important;
    color: #173f5f !important;
    border: 1px solid #b9cad8 !important;
    border-radius: 10px !important;
    font-weight: 750 !important;
}

[data-testid="stDownloadButton"] button:hover {
    background: #eef7f5 !important;
    border-color: #74c7bc !important;
    color: #0f766e !important;
}

/* ---------- Owner badge ---------- */
.ai-owner {
    position: fixed;
    left: 50%;
    bottom: 10px;
    transform: translateX(-50%);
    z-index: 99999;
    min-width: 175px;
    padding: 7px 13px;
    text-align: center;
    font-size: 10px;
    line-height: 1.35;
    color: #486581;
    background: rgba(255,255,255,.96);
    border: 1px solid #d9e2ec;
    border-radius: 10px;
    box-shadow: 0 5px 18px rgba(16,42,67,.10);
    pointer-events: none;
}

.ai-owner b {
    color: #173f5f;
    letter-spacing: .35px;
}

/* ---------- Responsive ---------- */
@media (max-width: 900px) {
    .block-container {
        padding-left: 1.1rem !important;
        padding-right: 1.1rem !important;
    }

    .ai-hero {
        padding: 28px 24px;
        border-radius: 18px;
    }

    .ai-hero h1 {
        font-size: 30px !important;
    }

    .ai-card {
        padding: 19px;
    }
}

@media (max-width: 640px) {
    .ai-hero h1 {
        font-size: 25px !important;
    }

    .ai-hero p {
        font-size: 14px !important;
    }

    .ai-badge {
        font-size: 11px;
    }
}


/* =========================================================
   FINAL VISIBILITY FIXES — ONLY FORM CONTROLS / CALCULATED MARKS
   No backend, generation, validation or document formatting changes.
   ========================================================= */

/* Make all sidebar controls clearly readable */
[data-testid="stSidebar"] [data-baseweb="select"] > div,
[data-testid="stSidebar"] [data-baseweb="input"] > div,
[data-testid="stSidebar"] [data-testid="stDateInput"] > div {
    background: #eef4f8 !important;
    border: 1px solid #7f9db2 !important;
    border-radius: 10px !important;
    box-shadow: none !important;
}

[data-testid="stSidebar"] [data-baseweb="select"]:hover > div,
[data-testid="stSidebar"] [data-testid="stDateInput"]:hover > div,
[data-testid="stSidebar"] [data-testid="stNumberInput"]:hover > div {
    border-color: #14b8a6 !important;
}

/* Strong, dark text inside select/date/number controls */
[data-testid="stSidebar"] [data-baseweb="select"] span,
[data-testid="stSidebar"] [data-baseweb="select"] input,
[data-testid="stSidebar"] [data-baseweb="input"] input,
[data-testid="stSidebar"] [data-testid="stDateInput"] input,
[data-testid="stSidebar"] [data-testid="stNumberInput"] input {
    color: #102a43 !important;
    -webkit-text-fill-color: #102a43 !important;
    font-weight: 700 !important;
}

/* Dropdown arrow area — no more white-on-white */
[data-testid="stSidebar"] [data-baseweb="select"] svg,
[data-testid="stSidebar"] [data-baseweb="select"] svg path {
    color: #0f766e !important;
    fill: #0f766e !important;
    stroke: #0f766e !important;
    opacity: 1 !important;
    visibility: visible !important;
}

/* Keep the dropdown arrow area visibly separated from the field. */
[data-testid="stSidebar"] [data-baseweb="select"] > div > div:last-child {
    background: #d7ebe8 !important;
    border-left: 1px solid #8fc5bf !important;
    border-radius: 0 9px 9px 0 !important;
    opacity: 1 !important;
    visibility: visible !important;
}

[data-testid="stSidebar"] [data-baseweb="select"] [role="combobox"] {
    color: #102a43 !important;
}

/* Date picker icon / right control */
[data-testid="stSidebar"] [data-testid="stDateInput"] svg {
    color: #0f766e !important;
    fill: #0f766e !important;
}

[data-testid="stSidebar"] [data-testid="stDateInput"] button {
    background: #d7ebe8 !important;
    color: #0f766e !important;
    border-left: 1px solid #a9ccc7 !important;
}

/* Number input: remove the extra outer white box; keep only the actual input row. */
[data-testid="stSidebar"] [data-testid="stNumberInput"] {
    background: transparent !important;
    border: 0 !important;
    box-shadow: none !important;
}

[data-testid="stSidebar"] [data-testid="stNumberInput"] > div {
    background: transparent !important;
    border: 0 !important;
    box-shadow: none !important;
}

[data-testid="stSidebar"] [data-testid="stNumberInput"] [data-baseweb="input"] > div {
    background: #eef4f8 !important;
    border: 1px solid #7f9db2 !important;
    border-radius: 10px !important;
    box-shadow: none !important;
}

/* Number input: clear dark minus/plus controls */
[data-testid="stSidebar"] [data-testid="stNumberInput"] button {
    background: #d7ebe8 !important;
    color: #102a43 !important;
    border-left: 1px solid #a9ccc7 !important;
    font-weight: 900 !important;
    font-size: 18px !important;
    opacity: 1 !important;
}

[data-testid="stSidebar"] [data-testid="stNumberInput"] button:hover {
    background: #bfe0dc !important;
    color: #0b625c !important;
}

[data-testid="stSidebar"] [data-testid="stNumberInput"] button svg {
    color: #102a43 !important;
    fill: #102a43 !important;
    stroke: #102a43 !important;
}

/* Choice / dropdown popup — readable options and highlighted selection */
[data-baseweb="popover"] [data-baseweb="menu"] {
    background: #ffffff !important;
    border: 1px solid #9bb4c6 !important;
    box-shadow: 0 12px 30px rgba(16,42,67,.22) !important;
}

[data-baseweb="popover"] [data-baseweb="option"] {
    color: #102a43 !important;
    background: #ffffff !important;
    font-weight: 600 !important;
}

[data-baseweb="popover"] [data-baseweb="option"]:hover,
[data-baseweb="popover"] [aria-selected="true"] {
    color: #0b625c !important;
    background: #e4f4f1 !important;
}

/* Calculated Marks + validation/error messages: dark, high contrast */
[data-testid="stSidebar"] .stAlert {
    background: #eef4f8 !important;
    border: 1px solid #8aa7ba !important;
    color: #102a43 !important;
    border-radius: 10px !important;
}

[data-testid="stSidebar"] .stAlert p,
[data-testid="stSidebar"] .stAlert span,
[data-testid="stSidebar"] .stAlert div {
    color: #102a43 !important;
    -webkit-text-fill-color: #102a43 !important;
    font-weight: 650 !important;
}

/* Calculated Marks label itself */
[data-testid="stSidebar"] .stMarkdown strong {
    color: #ffffff !important;
}

</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="ai-hero">
  <h1>Artificial Intelligence Based Smart Examination Paper Generator</h1>
  <p>Generate a structured college examination paper and answer key from your uploaded study material.</p>
  <span class="ai-badge">🔐 Private session • Multiple users can use the app simultaneously</span>
</div>
""", unsafe_allow_html=True)

# Keep the configuration in the left menu, styled as the final dark control panel.
with st.sidebar:
    st.markdown("## 🤖 AI Smart Exam Generator")
    st.caption("QUESTION CONFIGURATION")
    st.divider()
    st.markdown("### Exam Details")
    college_name=st.text_input('College Name',value='EXCEL COLLEGE FOR COMMERCE AND SCIENCE')
    course_name=st.text_input('Course / Degree',value='BACHELOR OF COMPUTER APPLICATIONS')
    subject_name=st.text_input('Subject Name',value='')
    subject_code=st.text_input('Subject Code',value='')
    semester=st.text_input('Semester',value='SEM-V')
    exam_session=st.text_input('Exam Session / Title',value='INTERNAL ASSESSMENT EXAMINATION (IAE) - I - JULY - 2026')
    set_name=st.selectbox('Question Paper Set',['SET-I','SET-II'])
    exam_date=st.date_input('Date',value=date.today())
    exam_time=st.text_input('Time / Duration',value='3 Hours')
    college_location=st.text_input('College Location',value='Pallakapalayam- 637 303, Namakkal Dt., Tamil Nadu')
    language_mode=st.radio('Question Paper Language',['English','Tamil','English + Tamil'],index=0)
    total_marks=st.number_input('Total Marks',min_value=1,value=50,step=1)
    include_signature=st.checkbox('Include Signature / Approval',value=False,help='Tick to add the approval table only at the end of the generated question paper.')

    st.markdown("### ＋ Add Part / Section")
    if 'sections_ui' not in st.session_state:
        st.session_state.sections_ui=[{'name':'Part A','question_type':'MCQ','questions_to_generate':10,'questions_to_attempt':10,'marks_per_question':1,'choice_type':'Answer All','or_choice':False,'passage_questions':5}]
    if st.button('＋ Add Part / Section',use_container_width=True):
        n=len(st.session_state.sections_ui)+1
        st.session_state.sections_ui.append({'name':f'Part {chr(64+n)}','question_type':'MCQ','questions_to_generate':1,'questions_to_attempt':1,'marks_per_question':1,'choice_type':'Answer All','or_choice':False,'passage_questions':5})
        st.rerun()
    remove_idx=None
    types=['MCQ','Paragraph / Descriptive']
    for i,sec in enumerate(st.session_state.sections_ui):
        with st.container(border=True):
            sec['name']=st.text_input(f'Section {i+1} Name',sec['name'],key=f'secname{i}')
            sec['question_type']=st.selectbox('Question Type',types,index=types.index(sec['question_type']) if sec['question_type'] in types else 0,key=f'sectype{i}')
            if sec['question_type']=='Paragraph / Descriptive':
                sec['or_choice']=st.checkbox('Either / OR Choice',value=bool(sec.get('or_choice',False)),key=f'orchoice{i}')
            else:
                sec['or_choice']=False
            sec['questions_to_generate']=st.number_input('Questions to Generate',min_value=1,value=int(sec['questions_to_generate']),key=f'gen{i}')
            sec['questions_to_attempt']=st.number_input('Questions to Attempt',min_value=1,value=int(sec['questions_to_attempt']),key=f'att{i}')
            sec['marks_per_question']=st.number_input('Marks per Question',min_value=1,value=int(sec['marks_per_question']),key=f'mark{i}')
            sec['choice_type']=st.selectbox('Choice', ['Answer All','Answer Any'], index=0 if sec['choice_type']=='Answer All' else 1,key=f'choice{i}')
            if len(st.session_state.sections_ui)>1 and st.button('Remove this section',key=f'remove{i}'): remove_idx=i
    if remove_idx is not None:
        st.session_state.sections_ui.pop(remove_idx); st.rerun()

    meta={'college_name':college_name.strip(),'course_name':course_name.strip(),'subject_name':subject_name.strip(),'subject_code':subject_code.strip(),'semester':semester.strip(),'exam_session':exam_session.strip(),'set_name':set_name,'exam_date':exam_date.strftime('%d-%m-%Y'),'exam_time':exam_time.strip(),'college_location':college_location.strip(),'language_mode':language_mode,'total_marks':int(total_marks),'header_translations':st.session_state.get('header_translations',{}),'sections':st.session_state.sections_ui,'include_signature':bool(include_signature)}
    errors,calc=validate_pattern(meta)
    st.markdown(f'**Calculated Marks:** {calc} / {int(total_marks)}')
    if errors:
        for e in errors: st.error(e)
    else:
        st.success('Marks pattern is valid.')

if not subject_name.strip() or not subject_code.strip():
    st.info('Enter Subject Name and Subject Code, then upload study material.')

st.markdown('<div class="step-title">1. Upload Study Material</div>', unsafe_allow_html=True)
st.markdown('<div class="step-sub">Upload one Unit / Lesson as PDF or Word document.</div>', unsafe_allow_html=True)
uploaded_files=st.file_uploader('Browse File',type=['pdf','docx'],accept_multiple_files=True)
if uploaded_files:
    st.success(f'{len(uploaded_files)} study-material file(s) uploaded.')

st.markdown('<div class="step-title">2. Generate</div>', unsafe_allow_html=True)
st.markdown('<div class="step-sub">AI follows only the selected sections and marks pattern.</div>', unsafe_allow_html=True)
if uploaded_files:
    errors,_=validate_pattern(meta); ready=bool(subject_name.strip() and subject_code.strip() and not errors)
    if not ready: st.warning('Fix the exam details / marks pattern before generating.')
    elif st.button('✨ Generate Question Paper + Answer Key',type='primary',use_container_width=True):
        with st.spinner('AI is reading the study material and generating the exact selected pattern...'):
            try:
                result=generate_question_paper(meta,uploaded_files)
                meta['header_translations']=result.get('header_translations',{}) or {}
                st.session_state['header_translations']=meta['header_translations']
                st.session_state['generated_result']=result
                st.session_state['generated_meta']=meta
                st.success('Question paper and answer key generated successfully.')
            except Exception as exc:
                st.error(str(exc))
else:
    st.info('Upload your study material to enable generation.')

result = st.session_state.get("generated_result")
saved_meta = st.session_state.get("generated_meta")

if result and saved_meta:

    st.markdown(
        '<div style="height:14px"></div>',
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="ai-card">',
        unsafe_allow_html=True
    )

    st.markdown(
        '<div class="step-title">3. Generated Question Paper Preview</div>',
        unsafe_allow_html=True
    )

    # ---------------------------------------------------------
    # GENERATED QUESTION PAPER PREVIEW
    # ---------------------------------------------------------
    for sec in result.get("sections", []):

        with st.expander(
            sec.get("name", "Section"),
            expanded=True
        ):

            for q in sec.get("questions", []):

                st.markdown(
                    f"**{q.get('q_no', '')}. {q.get('question', '')}**"
                )

                # MCQ
                if sec.get("question_type") == "MCQ":

                    o = q.get("options", {}) or {}

                    st.write(
                        f"A) {o.get('A', '')}   "
                        f"B) {o.get('B', '')}   "
                        f"C) {o.get('C', '')}   "
                        f"D) {o.get('D', '')}"
                    )

                # Paragraph / Descriptive with OR
                elif sec.get("or_choice") and q.get("or_question"):

                    st.markdown("**OR**")

                    st.markdown(
                        f"**{q.get('q_no', '')}. "
                        f"{q['or_question'].get('question', '')}**"
                    )

    # ---------------------------------------------------------
    # CREATE OUTPUT DIRECTORY
    # ---------------------------------------------------------
    out_dir = Path("generated_output") / SESSION_ID
    out_dir.mkdir(parents=True, exist_ok=True)

    # ---------------------------------------------------------
    # OUTPUT FILE PATHS
    # ---------------------------------------------------------
    qp_path = out_dir / (
        f"{saved_meta['subject_code']}_Question_Paper.docx"
    )

    ak_path = out_dir / (
        f"{saved_meta['subject_code']}_Answer_Key.docx"
    )

    # ---------------------------------------------------------
    # GENERATE DOCX ONLY IF IT DOES NOT ALREADY EXIST
    # This avoids regenerating the DOCX on every Streamlit rerun.
    # ---------------------------------------------------------
    if not qp_path.exists():
        generate_question_docx(
            saved_meta,
            result,
            qp_path
        )

    if not ak_path.exists():
        generate_answer_key_docx(
            saved_meta,
            result,
            ak_path
        )

    # ---------------------------------------------------------
    # DOWNLOAD BUTTONS
    # ---------------------------------------------------------
    c1, c2 = st.columns(2)

    with c1:
        st.download_button(
            "📄 Download Question Paper",
            data=qp_path.read_bytes(),
            file_name=qp_path.name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    with c2:
        st.download_button(
            "📝 Download Answer Key",
            data=ak_path.read_bytes(),
            file_name=ak_path.name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    st.markdown(
        "</div>",
        unsafe_allow_html=True
    )


st.markdown(
    '<div class="ai-owner">'
    '<b>WEBSITE OWNERS</b><br>'
    'K.ARULRAJA<br>'
    'K.BOOPATHIRAJA'
    '</div>',
    unsafe_allow_html=True
)