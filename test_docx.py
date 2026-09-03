from docx import Document

def generate_docx(questions, filepath):
    doc = Document()

    doc.add_heading("XYZ College of Engineering", level=1)
    doc.add_paragraph("CAT-1 - Data Structures")
    doc.add_paragraph("")  # empty line spacing

    for q in questions:
        doc.add_paragraph(f"Q{q['q_no']}. {q['question']} [{q['marks']} marks]")

    doc.save(filepath)


sample_questions = [
    {"q_no": 1, "question": "Explain Floyd's Cycle Detection Algorithm to detect a loop in a linked list.", "marks": 10},
    {"q_no": 2, "question": "Write an algorithm to reverse a doubly linked list in-place.", "marks": 10},
]

generate_docx(sample_questions, "sample_question_paper.docx")
print("Word document generated successfully!")